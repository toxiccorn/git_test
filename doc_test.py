from docx import Document
from docx.shared import Cm, Pt
#from docx.enum.text import WD_ALIGN_PARAGRAPH

# Создаем файл ведомости
document = Document()
section = document.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1)

default_style = document.styles['Normal']
default_font = default_style.font
default_font.name = 'Times New Roman'
default_font.size = Pt(10)
default_font.bold = True

doc_rsreu_head = document.add_paragraph("ФГБОУ ВО «Рязанский"
                                        " государственный радиотехнический университет им. В.Ф. Уткина»")
doc_rsreu_head.alignment = 1
doc_rsreu_head.style = document.styles['Normal']

doc_name = document.add_paragraph('Экзаменационная ведомость')
doc_name.alignment = 1
doc_name.style = document.styles['Normal']

doc_exam_name = document.add_paragraph('Имя экзамена')
doc_exam_name.alignment = 1
doc_exam_name.style = document.styles['Normal']

doc_exam_date = document.add_paragraph("Дата: " + "случайная дата")
doc_exam_date.alignment = 1
doc_exam_date.style = document.styles['Normal']


document.save('test.docx')