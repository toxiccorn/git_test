import csv
import os
import vars

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from Student import Student

# Поиск и выбор файла
file_list = []
for file in os.listdir():
    if file.endswith('.csv'):
        file_list.append(file)

for index, file in enumerate(file_list):
    print(str(index) + ' ' + file)

user_file_pick = int(input('выберите номер файла'))
filename = file_list[user_file_pick]

temp_file = 'temp_' + filename
encode = 'utf-8'

# Ищем тест и его параметры
for index, subject in enumerate(vars.subjects):
    print(str(index) + ' ' + subject['name'])
subject_index = int(input('выберете предмет'))
subject = vars.subjects[subject_index]

subject_name = subject['name']
min_score = subject['min_score']
question_complete_mark = subject['question_complete_mark']
a_scores_table = subject['custom_a_table']
b_scores_table = subject['custom_b_table']
a_start = subject['a_questions_start_col']
a_end = subject['a_questions_end_col']
b_start = subject['b_questions_start_col']
b_end = subject['b_questions_end_col']
test_date_col = subject['test_date_col']
b_questions = len(subject['custom_b_table'].items()) - 1

date_user_string = input('Введите дату экзамена в формате dd.mm.yyyy')

# Перевариваем дату для поиска по таблице
date_splited_strings = date_user_string.split('.')
test_day = date_splited_strings[0]
test_day_int = int(test_day)
test_month = date_splited_strings[1]
test_month_search_string = vars.months[int(test_month) - 1]
test_year = date_splited_strings[2]
search_date_string = str(test_day_int) + ' ' + test_month_search_string + ' ' + test_year

# Открывает файл и чичстит его от лишних кавычек
with open(filename, encoding=encode) as file:
    data = file.readlines()
    del data[-1]
    for line in data:
        line.replace('"', '')

# Пишет во временный файл читаемые данные
with open(temp_file, 'w', encoding=encode) as tfile:
    for line in data:
        tfile.write(line)

# Открывает временный файл
with open(temp_file, 'r', encoding=encode) as tfile:
    students = []
    csvdata = csv.reader(tfile)
    header_row = next(csvdata)

# Cчитывает данные о необходимых студентах
    for row in csvdata:
        if search_date_string in row[test_date_col]:
            current_student = Student(row[0] + ' ' + row[1], row[9])

            for a_list in range(a_start, a_end + 1):
                current_student.a_marks.append(row[a_list])
            for b_list in range(b_start, b_end + 1):
                current_student.b_marks.append(row[b_list])

            current_student.clear_marks()
            current_student.calculate_primary_score(question_complete_mark)
            current_student.calculate_secondary_score(a_scores_table, b_scores_table)

            if current_student.sum_secondary_score < min_score:
                result = current_student.autoscore(min_score, b_questions, question_complete_mark, a_scores_table,
                                                   b_scores_table)
                if result == 'success':
                    print(current_student.full_name + ' autocorrected ' + str(current_student.autocorrect_attemps) +
                          ' times')
                else:
                    print(current_student.full_name + ' is dibil')

            students.append(current_student)

# Вывод информации в консоль
for student in students:
    print(student.full_name + ' - ' + str(student.sum_secondary_score))

# Удаляет временный файл
os.remove(temp_file)

# Создаем файл ведомости
document = Document()
# Устанавливаем поля
sections = document.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)

# Определяем стиль по умолчанию
object_style = document.styles
paragraph_style = object_style.add_style("Paragraph style", WD_STYLE_TYPE.PARAGRAPH)
paragraph_font = paragraph_style.font
paragraph_font.size = Pt(12)
paragraph_font.name = 'Times New Roman'

doc_rsreu_head = document.add_paragraph(style='Paragraph style')
doc_rsreu_head.alignment = 1
doc_rsreu_head.add_run("ФГБОУ ВО «Рязанский государственный радиотехнический университет им. В.Ф. Уткина»").bold = True

doc_name = document.add_paragraph(style='Paragraph style')
doc_name.alignment = 1
doc_name.add_run('Экзаменационная ведомость').bold = True

doc_exam_date = document.add_paragraph(style='Paragraph style')
doc_exam_date.alignment = 1
doc_exam_date.add_run("Дата: " + date_user_string).bold = True

doc_exam_name = document.add_paragraph(style='Paragraph style')
doc_exam_name.alignment = 1
doc_exam_name.add_run(subject_name.title()).bold = True


# Таблица
table = document.add_table(cols=3, rows=1)
table.style = "Table Grid"
table.heigh_rule = WD_ROW_HEIGHT_RULE.EXACTLY
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

heading_row = table.rows[0].cells
heading_row[0].text = '№ п/п'
heading_row[0].paragraphs[0].alignment = 1
heading_row[1].text = 'ФИО'
heading_row[1].paragraphs[0].alignment = 1
heading_row[2].text = 'БАЛЛ'
heading_row[2].paragraphs[0].alignment = 1

for index, student in enumerate(students):
    new_row = table.add_row()
    number = new_row.cells[0].paragraphs[0]
    number.style = "Paragraph style"
    number.alignment = 1
    number.add_run(str(index))

    name = new_row.cells[1].paragraphs[0]
    name.style = "Paragraph style"
    name.alignment = 0
    name.add_run(student.full_name)

    score = new_row.cells[2].paragraphs[0]
    score.alignment = 1
    score.style = "Paragraph style"
    score.add_run(str(student.sum_secondary_score))

for row in table.rows:
    row.height = Cm(1.5)
    for index, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if index != 1:
            cell.width = Cm(2)
        else:
            cell.width = Cm(14)
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.bold = True

document.save('demo.docx')
